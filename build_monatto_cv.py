from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION

OUT = "Juan_Manuel_Gomez_Palma_CV_Technical_Integrations_Analyst.docx"

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(0.78)
section.right_margin = Inches(0.78)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles

def set_font(style, name="Calibri", size=10, bold=False, color=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor(*color)

normal = styles["Normal"]
set_font(normal, size=9.6)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.0

for style_name in ("List Bullet",):
    s = styles[style_name]
    set_font(s, size=9.4)
    s.paragraph_format.left_indent = Inches(0.22)
    s.paragraph_format.first_line_indent = Inches(-0.14)
    s.paragraph_format.space_after = Pt(1)
    s.paragraph_format.line_spacing = 1.0

def add_style(name, size, bold=False, color=(46, 116, 181), before=0, after=0):
    s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    set_font(s, size=size, bold=bold, color=color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.line_spacing = 1.0
    s.paragraph_format.keep_with_next = True
    return s

name_style = add_style("CV Name", 17, True, (11, 37, 69), 0, 1)
headline_style = add_style("CV Headline", 10.5, False, (31, 77, 120), 0, 1)
contact_style = add_style("CV Contact", 9.2, False, (0, 0, 0), 0, 5)
section_style = add_style("CV Section", 11.5, True, (46, 116, 181), 7, 2)
summary_style = add_style("CV Summary", 9.6, False, (0, 0, 0), 0, 3)
competency_style = add_style("CV Competencies", 9.3, False, (0, 0, 0), 0, 2)
company_style = add_style("CV Company", 10.2, True, (11, 37, 69), 2, 0)
meta_style = add_style("CV Meta", 9.2, False, (0, 0, 0), 0, 0)
tech_style = add_style("CV Tech", 8.9, False, (80, 80, 80), 0, 1)
project_style = add_style("CV Project", 9.8, True, (11, 37, 69), 2, 0)
project_desc_style = add_style("CV Project Desc", 9.2, False, (0, 0, 0), 0, 0)
skill_label_style = add_style("CV Skill Label", 9.2, True, (0, 0, 0), 1, 0)
education_style = add_style("CV Education", 9.2, False, (0, 0, 0), 1, 0)

def add_p(text, style, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p

def add_section(title):
    add_p(title, section_style)

def add_experience(company, role_dates, technologies, bullets):
    add_p(company, company_style)
    add_p(role_dates, meta_style)
    add_p(technologies, tech_style)
    for bullet in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(bullet)

def add_project(name, description, technologies):
    add_p(name, project_style)
    add_p(description, project_desc_style)
    add_p(technologies, tech_style)

# Header
add_p("Juan Manuel Gomez Palma", name_style)
add_p("Analista de Integraciones Técnicas | APIs, SaaS y Support Operations", headline_style)
add_p("juanmgomezpalma@gmail.com | +52 98 55 93 7696 | Yucatán, México | linkedin.com/in/juanmanuel-gomezpalma | github.com/juanmgomezp", contact_style)

add_section("Professional Summary")
add_p("Ingeniero en Sistemas Computacionales con experiencia desde 2020 en implementación de software empresarial, soporte técnico B2B y administración SaaS. Especializado en validar flujos de APIs REST, diagnosticar aplicaciones mediante solicitudes HTTP y evidencia técnica, y colaborar con Desarrollo durante pruebas y liberaciones. Combina SQL Server, Postman, documentación técnica y análisis de causa raíz para reducir fricción operativa y fortalecer la experiencia del cliente.", summary_style)

add_section("Core Competencies")
add_p("Integraciones de sistemas | APIs REST | Validación con Postman | SQL Server | Diagnóstico técnico | Documentación técnica | Administración SaaS | Análisis de causa raíz", competency_style)

add_section("Professional Experience")
add_experience(
    "Ekasoft",
    "Support Administrator | Abril 2026 – Actualidad",
    "Zoho Desk | Zoho SalesIQ | Zoho Deluge | SQL Server | Python",
    [
        "Administra licencias, proyectos, productos y usuarios del entorno SaaS, y valida configuraciones antes de liberar nuevos ambientes.",
        "Valida diagnósticos técnicos antes de escalamientos a Desarrollo y coordina la resolución de incidentes críticos para más de 100 clientes empresariales.",
        "Automatiza el seguimiento de escalamientos en Zoho Desk para medir tiempos de atención, identificar cuellos de botella y analizar su impacto operativo.",
    ],
)

add_experience(
    "Ekasoft",
    "Support Lead | Abril 2025 – Marzo 2026",
    "SQL Server | Browser Developer Tools | HTTP | REST APIs | JSON | Zoho Desk",
    [
        "Lideró la migración de aproximadamente 20 clientes empresariales hacia la plataforma SaaS sin pérdida de información ni interrupciones relevantes del servicio.",
        "Analizó errores de consola y solicitudes de red para diagnosticar incidencias SaaS, reunir evidencia técnica y coordinar investigaciones con Desarrollo y TI de clientes.",
        "Validó soluciones antes de su liberación al cliente y coordinó un equipo de cinco ingenieros durante incidentes de alta complejidad.",
    ],
)

add_experience(
    "Ekasoft",
    "Senior Support Engineer | Agosto 2022 – Marzo 2025",
    "SQL Server | REST APIs | Postman | HTTP | JSON | C# | .NET Core",
    [
        "Validó funcionalmente APIs REST antes de liberaciones a producción, verificando autenticación, endpoints, comportamiento funcional y tiempos de respuesta con Postman.",
        "Analizó solicitudes HTTP, respuestas del servidor y errores de navegador para investigar incidencias en aplicaciones empresariales de escritorio y SaaS.",
        "Comparó el comportamiento de aplicaciones con procedimientos almacenados para aportar evidencia a investigaciones de rendimiento con Desarrollo.",
    ],
)

doc.add_page_break()

add_experience(
    "Ekasoft",
    "Support Engineer | Febrero 2022 – Julio 2022",
    "SQL Server | SQL Profiler | SQL Server Management Studio | Freshdesk | Microsoft Teams",
    [
        "Diagnosticó incidencias de clientes mediante SQL Server y SQL Profiler, identificando causas raíz y escalando casos a Desarrollo cuando fue necesario.",
        "Documentó casos complejos y elaboró guías para configuraciones y procesos operativos frecuentes, facilitando investigaciones posteriores.",
        "Resolvió más del 60% de los incidentes durante la primera hora desde la creación del ticket.",
    ],
)

add_experience(
    "Ekasoft",
    "Implementation Engineer | Octubre 2020 – Enero 2022",
    "SQL Server | Consoft RAD Platform | Microsoft Excel | Microsoft Word | Microsoft Teams",
    [
        "Acompañó implementaciones de software empresarial desde la instalación y configuración hasta la validación funcional, capacitación y adopción del cliente.",
        "Documentó configuraciones y actividades técnicas, y diseñó un formato de memoria técnica adoptado como estándar por el área de Implementación.",
        "Participó en aproximadamente 30 órdenes de servicio entre implementaciones y renovaciones, con una tasa de adopción cercana al 99% en los proyectos asignados.",
    ],
)

add_section("Selected Projects")
add_project(
    "Validación Funcional de API REST para Consulta de Comprobantes con Carta Porte",
    "Validó autenticación, endpoints, respuestas y tiempos de ejecución en distintos ambientes; aportó evidencia para investigar rendimiento y confirmó el funcionamiento posterior a las optimizaciones.",
    "REST APIs | HTTP | JSON | Postman | SQL Server | Stored Procedures",
)
add_project(
    "Migración y Adopción de Clientes hacia Plataforma SaaS",
    "Preparó, configuró y validó entornos SaaS, coordinó incidencias con Desarrollo y acompañó la adopción de aproximadamente 20 clientes sin pérdida de información.",
    "Plataforma SaaS de Ekasoft | SQL Server | SQL Server Management Studio | Microsoft Teams",
)

add_section("Technical Skills")
add_p("APIs e integración: REST APIs | HTTP | JSON | Bearer Token Authentication | Postman", skill_label_style)
add_p("Datos y programación: SQL Server | Stored Procedures | Python | C# | .NET Core", skill_label_style)
add_p("Plataformas y herramientas: SaaS Administration | Zoho Desk | Zoho SalesIQ | Zoho Deluge | Browser Developer Tools", skill_label_style)

add_section("Education")
add_p("Ingeniero en Sistemas Computacionales | Instituto Tecnológico Superior de Valladolid | 2014", education_style)

add_section("Languages")
add_p("Español — Nativo | Inglés — B2", education_style)

# Make line-based metadata stay with the following content where possible.
for p in doc.paragraphs:
    p.paragraph_format.widow_control = True

doc.core_properties.title = "CV - Technical Integrations Analyst"
doc.core_properties.author = "Juan Manuel Gomez Palma"
doc.save(OUT)
print(OUT)
