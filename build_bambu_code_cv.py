from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


OUT = "Juan_Manuel_Gomez_Palma_CV_Lider_Soporte_Clientes_Bambu_Code.docx"

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.72)
section.right_margin = Inches(0.72)

styles = doc.styles


def set_font(style, size, bold=False, color=None):
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor(*color)


def add_style(name, size, bold=False, color=(0, 0, 0), before=0, after=0):
    style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    set_font(style, size, bold, color)
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = 1.0
    return style


normal = styles["Normal"]
set_font(normal, 9.1)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.0

bullet_style = styles["List Bullet"]
set_font(bullet_style, 9.0)
bullet_style.paragraph_format.left_indent = Inches(0.22)
bullet_style.paragraph_format.first_line_indent = Inches(-0.13)
bullet_style.paragraph_format.space_after = Pt(1)
bullet_style.paragraph_format.line_spacing = 1.0

name_style = add_style("CV Name", 17, True, (11, 37, 69), 0, 0)
headline_style = add_style("CV Headline", 10.2, False, (31, 77, 120), 0, 1)
contact_style = add_style("CV Contact", 8.8, False, (0, 0, 0), 0, 4)
section_style = add_style("CV Section", 11.2, True, (46, 116, 181), 5, 1)
summary_style = add_style("CV Summary", 9.1, False, (0, 0, 0), 0, 2)
competencies_style = add_style("CV Competencies", 8.9, False, (0, 0, 0), 0, 1)
company_style = add_style("CV Company", 9.8, True, (11, 37, 69), 2, 0)
meta_style = add_style("CV Meta", 8.9, False, (0, 0, 0), 0, 0)
tech_style = add_style("CV Tech", 8.5, False, (89, 89, 89), 0, 1)
project_style = add_style("CV Project", 9.2, True, (11, 37, 69), 2, 0)
project_desc_style = add_style("CV Project Description", 8.9, False, (0, 0, 0), 0, 0)
skill_style = add_style("CV Skill", 8.9, False, (0, 0, 0), 1, 1)


def add_paragraph(text, style):
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(text)
    paragraph.paragraph_format.widow_control = True
    return paragraph


def add_section(title):
    paragraph = add_paragraph(title, section_style)
    paragraph.paragraph_format.keep_with_next = True


def add_experience(company, role_dates, technologies, bullets):
    company_paragraph = add_paragraph(company, company_style)
    company_paragraph.paragraph_format.keep_with_next = True
    role_paragraph = add_paragraph(role_dates, meta_style)
    role_paragraph.paragraph_format.keep_with_next = True
    tech_paragraph = add_paragraph(technologies, tech_style)
    tech_paragraph.paragraph_format.keep_with_next = True
    for bullet in bullets:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(bullet)
        paragraph.paragraph_format.widow_control = True


def add_project(name, description, technologies):
    project_paragraph = add_paragraph(name, project_style)
    project_paragraph.paragraph_format.keep_with_next = True
    description_paragraph = add_paragraph(description, project_desc_style)
    description_paragraph.paragraph_format.keep_with_next = True
    add_paragraph(technologies, tech_style)


# Header
add_paragraph("Juan Manuel Gomez Palma", name_style)
add_paragraph("Líder de Soporte al Cliente | Support Operations | SaaS", headline_style)
add_paragraph(
    "juanmgomezpalma@gmail.com | +52 98 55 93 7696 | Yucatán, México | "
    "linkedin.com/in/juanmanuel-gomezpalma | github.com/juanmgomezp",
    contact_style,
)

add_section("Professional Summary")
add_paragraph(
    "Ingeniero en Sistemas Computacionales con trayectoria desde 2020 en soporte técnico, "
    "operaciones de soporte y administración de plataformas SaaS. Lidera equipos remotos, "
    "gestiona indicadores de servicio y desarrolla automatizaciones, autoservicio y análisis "
    "operativo para mejorar la experiencia del cliente. Experiencia con clientes empresariales "
    "B2B y usuarios finales de productos de facturación.",
    summary_style,
)

add_section("Core Competencies")
add_paragraph(
    "Liderazgo remoto | Coaching y feedback | KPIs y SLA | Support Operations | SaaS | "
    "Mejora continua | Customer Experience | Gestión de escalamientos | Automatización | IA aplicada",
    competencies_style,
)

add_section("Professional Experience")
add_experience(
    "Ekasoft",
    "Support Administrator | Abril 2026 – Actualidad",
    "Zoho Desk | Zoho SalesIQ | Zoho Deluge | SQL Server | Python | SQLite",
    [
        "Coordina de forma 100% remota un equipo de tres ingenieros, combinando seguimiento operativo, priorización y sesiones bimestrales de feedback.",
        "Analiza indicadores de todos los equipos de soporte, diseña métricas y entrega reportes ejecutivos para identificar riesgos y oportunidades de mejora.",
        "Automatiza en Zoho Desk el ciclo de escalamientos hacia Desarrollo para medir tiempos, cuellos de botella e impacto sobre los SLA.",
        "Redujo 60% las solicitudes por licencias vencidas mediante un proceso preventivo coordinado con Key Account Management.",
        "Rediseña el flujo de chatbot hacia autoservicio mediante base de conocimientos para reducir carga operativa y diferenciar casos resueltos de escalados.",
    ],
)

add_experience(
    "Ekasoft",
    "Support Lead | Abril 2025 – Marzo 2026",
    "SQL Server | Freshdesk | Zoho Desk | C# | .NET Core | REST APIs",
    [
        "Coordinó un equipo de cinco ingenieros, balanceando cargas, guiando incidentes críticos y validando diagnósticos antes de escalamientos a Desarrollo.",
        "Mantuvo el SLA de primer contacto por encima de 95%, el SLA de resolución por encima de 90% y la satisfacción en tickets cerrados por encima de 95%.",
        "Lideró la migración de más de 20 clientes empresariales a la plataforma SaaS, con 99% de disponibilidad y sin pérdida de información.",
        "Coordinó con Desarrollo, Comercial y Customer Success la atención de clientes estratégicos, migraciones y mejoras operativas.",
    ],
)

add_experience(
    "Ekasoft",
    "Senior Support Engineer | Agosto 2022 – Marzo 2025",
    "SQL Server | C# | .NET Core | Postman | REST APIs | Freshdesk",
    [
        "Coordinó técnicamente el equipo especializado en facturación electrónica y analítica fiscal, apoyando decisiones, diagnósticos y desarrollo del equipo.",
        "Mantuvo el SLA de primera respuesta por encima de 97%, el de resolución por encima de 94% e incrementó la respuesta a encuestas por encima de 40%.",
        "Redujo aproximadamente 10% del volumen mensual de tickets mediante mejoras al producto y soluciones internas orientadas a causas recurrentes.",
        "Validó una API REST empresarial antes de producción y aportó evidencia técnica para apoyar la optimización de rendimiento con Desarrollo.",
    ],
)

add_experience(
    "Ekasoft",
    "Support Engineer | Febrero 2022 – Julio 2022",
    "SQL Server | SQL Profiler | Freshdesk | SQL Server Management Studio",
    [
        "Diagnosticó incidentes de productos de facturación electrónica para clientes empresariales y usuarios finales, identificando causas raíz con SQL Server.",
        "Resolvió más de 60% de los incidentes durante la primera hora desde la creación del ticket y mantuvo SLA de primera respuesta y resolución por encima de 95%.",
        "Documentó casos complejos y guías operativas para facilitar investigaciones posteriores y atención consistente al cliente.",
    ],
)

add_experience(
    "Ekasoft",
    "Implementation Engineer | Octubre 2020 – Enero 2022",
    "SQL Server | Consoft RAD Platform | Microsoft Teams | Microsoft Excel",
    [
        "Acompañó implementaciones de software desde la configuración y validación funcional hasta la capacitación de usuarios finales y adopción del sistema.",
        "Diseñó un formato de memoria técnica adoptado como estándar por el área de Implementación para mejorar la consistencia documental.",
        "Participó en aproximadamente 30 órdenes de servicio entre implementaciones y renovaciones, con adopción cercana a 99% en los proyectos asignados.",
    ],
)

add_section("Selected Projects")
add_project(
    "Dashboard de Métricas para Escalamientos a Desarrollo",
    "Diseñó la automatización y el dashboard que reconstruyen el ciclo de escalamiento, miden su efecto en SLA y permiten priorizar mejoras de producto con datos operativos.",
    "Zoho Desk | Zoho Deluge | Python | SQLite | IA asistida",
)
add_project(
    "Rediseño del Proceso de Atención mediante Chatbot",
    "Diseñó un flujo de autoservicio que usa la base de conocimientos, crea tickets solo cuando se requiere un agente y mide casos resueltos frente a escalados.",
    "Zoho SalesIQ | Zoho Desk | Zoho Deluge | IA asistida",
)
add_project(
    "Plataforma Centralizada de Encuestas y Analítica Comercial",
    "Automatizó la distribución y análisis de encuestas de satisfacción y NPS, centralizando resultados para apoyar decisiones sobre experiencia del cliente.",
    "SQL Server | C# | ASP.NET MVC | JavaScript",
)

add_section("Technical Skills")
add_paragraph("Support Operations: Gestión de SLA | Escalation Management | Root Cause Analysis | KPI Design | Customer Experience", skill_style)
add_paragraph("Automatización y datos: SQL Server | Stored Procedures | Zoho Deluge | Python | SQLite | Operational Analytics", skill_style)
add_paragraph("Plataformas y desarrollo: SaaS Administration | Zoho Desk | Zoho SalesIQ | C# | .NET Core | REST APIs | Postman", skill_style)

add_section("Education")
add_paragraph("Instituto Tecnológico Superior de Valladolid | Ingeniero en Sistemas Computacionales | 2014", skill_style)

add_section("Languages")
add_paragraph("Español — Nativo | Inglés — B2", skill_style)

doc.core_properties.title = "CV - Líder de Soporte al Cliente"
doc.core_properties.author = "Juan Manuel Gomez Palma"
doc.save(OUT)
print(OUT)
